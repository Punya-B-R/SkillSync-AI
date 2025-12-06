"""
Resume parsing service for extracting information from PDF, DOCX, and TXT files.
"""
import logging
import re
import io
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)

class ResumeParser:
    """Parse resume files and extract structured data."""
    
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB in bytes
    ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt'}
    
    def __init__(self):
        pass
    
    def _clean_text(self, text):
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw text string
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _count_words(self, text):
        """
        Count words in text.
        
        Args:
            text: Text string
            
        Returns:
            int: Word count
        """
        if not text:
            return 0
        words = text.split()
        return len(words)
    
    def parse_pdf(self, file):
        """
        Parse PDF resume file.
        
        Args:
            file: File object (BytesIO or file-like object)
            
        Returns:
            dict: Parsed resume data with raw_text, word_count, file_name, file_type
        """
        try:
            logger.info("Starting PDF parsing")
            file.seek(0)  # Reset file pointer
            pdf_reader = PdfReader(file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    logger.debug(f"Extracted text from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                    continue
            
            raw_text = '\n'.join(text_parts)
            
            if not raw_text or not raw_text.strip():
                raise ValueError("PDF file appears to be empty or contains no extractable text")
            
            cleaned_text = self._clean_text(raw_text)
            word_count = self._count_words(cleaned_text)
            
            logger.info(f"PDF parsed successfully. Word count: {word_count}")
            
            return {
                'raw_text': cleaned_text,
                'word_count': word_count,
                'file_type': 'pdf'
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF file: {str(e)}")
    
    def parse_docx(self, file):
        """
        Parse DOCX resume file.
        
        Args:
            file: File object (BytesIO or file-like object)
            
        Returns:
            dict: Parsed resume data with raw_text, word_count, file_name, file_type
        """
        try:
            logger.info("Starting DOCX parsing")
            file.seek(0)  # Reset file pointer
            doc = Document(file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            raw_text = '\n'.join(text_parts)
            
            if not raw_text or not raw_text.strip():
                raise ValueError("DOCX file appears to be empty or contains no text")
            
            cleaned_text = self._clean_text(raw_text)
            word_count = self._count_words(cleaned_text)
            
            logger.info(f"DOCX parsed successfully. Word count: {word_count}")
            
            return {
                'raw_text': cleaned_text,
                'word_count': word_count,
                'file_type': 'docx'
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
    
    def parse_txt(self, file):
        """
        Parse TXT resume file.
        
        Args:
            file: File object (BytesIO or file-like object)
            
        Returns:
            dict: Parsed resume data with raw_text, word_count, file_name, file_type
        """
        try:
            logger.info("Starting TXT parsing")
            file.seek(0)  # Reset file pointer
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            raw_text = None
            
            for encoding in encodings:
                try:
                    file.seek(0)
                    raw_text = file.read().decode(encoding)
                    logger.debug(f"Successfully decoded with {encoding}")
                    break
                except (UnicodeDecodeError, AttributeError):
                    continue
            
            if raw_text is None:
                # If all encodings fail, try reading as bytes and replacing errors
                file.seek(0)
                raw_text = file.read().decode('utf-8', errors='replace')
                logger.warning("Used UTF-8 with error replacement for decoding")
            
            if not raw_text or not raw_text.strip():
                raise ValueError("TXT file appears to be empty")
            
            cleaned_text = self._clean_text(raw_text)
            word_count = self._count_words(cleaned_text)
            
            logger.info(f"TXT parsed successfully. Word count: {word_count}")
            
            return {
                'raw_text': cleaned_text,
                'word_count': word_count,
                'file_type': 'txt'
            }
            
        except Exception as e:
            logger.error(f"Error parsing TXT: {str(e)}")
            raise ValueError(f"Failed to parse TXT file: {str(e)}")
    
    def parse_resume(self, file, filename):
        """
        Parse resume file based on type.
        
        Args:
            file: File object from request
            filename: Original filename
            
        Returns:
            dict: Parsed resume data with raw_text, word_count, file_name, file_type
            
        Raises:
            ValueError: If file type is invalid, file is corrupted, or empty
        """
        try:
            logger.info(f"Parsing resume: {filename}")
            
            # Validate file type
            file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
            if file_ext not in self.ALLOWED_EXTENSIONS:
                raise ValueError(f"Invalid file type. Allowed types: {', '.join(self.ALLOWED_EXTENSIONS)}")
            
            # Check file size
            file.seek(0, 2)  # Seek to end
            file_size = file.tell()
            file.seek(0)  # Reset to beginning
            
            if file_size > self.MAX_FILE_SIZE:
                raise ValueError(f"File size ({file_size / 1024 / 1024:.2f}MB) exceeds maximum allowed size (5MB)")
            
            if file_size == 0:
                raise ValueError("File is empty")
            
            # Convert file to BytesIO if needed for parsing
            if not hasattr(file, 'read'):
                raise ValueError("Invalid file object")
            
            # Parse based on file type
            if file_ext == 'pdf':
                result = self.parse_pdf(file)
            elif file_ext == 'docx':
                result = self.parse_docx(file)
            elif file_ext == 'txt':
                result = self.parse_txt(file)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Add filename to result
            result['file_name'] = filename
            
            logger.info(f"Resume parsed successfully: {filename}")
            return result
            
        except ValueError as e:
            logger.error(f"Validation error: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error parsing resume: {str(e)}")
            raise ValueError(f"Failed to parse resume file: {str(e)}")

